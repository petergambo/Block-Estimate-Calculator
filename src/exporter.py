from docx import Document
from docx.shared import Inches
import os

def export_to_word(title, filename, import_folder, export_folder, records, total_blocks):
    dir_path = os.path.dirname(os.path.realpath(__file__))

    document = Document()

    document.add_heading(title, 0)

    p = document.add_paragraph('Block Estimate for ')
    p.add_run('Client').bold = True
    p.add_run(' prepared by ')
    p.add_run('Peebot BOQ Analzer.').italic = True

    document.add_picture(os.path.join(import_folder, 'preview.png'), width=Inches(4.25))

    records = records

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Wallname'
    hdr_cells[1].text = 'Block(s)'

    for wallname, block in records:
        row_cells = table.add_row().cells
        row_cells[0].text = "Wall " + str(wallname)
        row_cells[1].text = str(block)
    row_cells[0].text = "SUB-TOTAL"
    row_cells[1].text = str(total_blocks)

    document.add_page_break()

    fullfilename = filename + ".docx"

    try:

        document.save(os.path.join(export_folder, fullfilename))
        status = "\t\tDocument {} created successfully".format(fullfilename)

    except:
        status = "\t\tDocument creation Failed! Check document path"

    return status